#!/usr/bin/env python3
"""Manuscript markdown -> docx (BMC submission format).

Fixed layout:
- Markdown paragraphs (blocks separated by blank lines) are merged into ONE
  Word paragraph (source soft line-wraps do not create separate paragraphs).
- Headings use built-in styles with consistent fonts.
- Tables get explicit column widths and a clean style.
- Figures are centered at a fixed width before their legends.
- References render as hanging-indent paragraphs.

Writes paper/output/doc/manuscript.docx (or given output).
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

PROJ = Path(__file__).resolve().parents[2]
MD = PROJ / "paper" / "output" / "doc" / "manuscript.md"
FIG_DIR = PROJ / "paper" / "figures"

FIG_FILES = {
    1: "fig1.png", 2: "fig2.png", 3: "fig3.png", 4: "fig4.png",
    5: "fig5.png", 6: "fig6.png", 7: "fig7.png", 8: "fig8.png",
    9: "fig9_subgroup_forest.png", 10: "fig10_error_cases.png",
}

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
R = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"
# internal hyperlink relationship id (anchor target)
INTERNAL_LINK_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/internalLink"


def add_internal_hyperlink(paragraph, text, anchor):
    """Add an internal hyperlink (bookmark anchor) as a run in paragraph.

    Word internal links use w:hyperlink with an 'anchor' attribute pointing to a
    bookmark name (no relationship needed), so Ctrl+click jumps to the target.
    """
    hyperlink = paragraph._element.makeelement(
        W + "hyperlink", {W + "anchor": anchor, W + "history": "1"})
    new_run = paragraph._element.makeelement(W + "r", {})
    rPr = paragraph._element.makeelement(W + "rPr", {})
    color = paragraph._element.makeelement(W + "color", {W + "val": "0563C1"})
    rPr.append(color)
    u = paragraph._element.makeelement(W + "u", {W + "val": "single"})
    rPr.append(u)
    new_run.append(rPr)
    t = paragraph._element.makeelement(W + "t", {})
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)


_bookmark_seq = [0]


def _next_bm_id():
    _bookmark_seq[0] += 1
    return str(_bookmark_seq[0])


def add_bookmark(paragraph, name):
    """Insert a bookmark that WRAPS the paragraph text so Word recognizes it.

    bookmarkStart goes before the first run, bookmarkEnd after the last run;
    a bookmark spanning at least one run is a valid, Word-recognized anchor.
    """
    start = paragraph._element.makeelement(W + "bookmarkStart", {
        W + "id": _next_bm_id(),
        W + "name": name,
    })
    end = paragraph._element.makeelement(W + "bookmarkEnd", {W + "id": str(_bookmark_seq[0])})
    paragraph._element.insert(0, start)
    paragraph._element.append(end)


def add_rich_text(par, text, enable_links=False):
    """Add text with **bold** markers rendered as bold runs, and [n] citations
    as internal hyperlinks when enable_links is True."""
    # split into citation groups [..] and normal text
    tokens = re.split(r"(\[[0-9, ]+\])", text)
    for tok in tokens:
        if not tok:
            continue
        cm = re.match(r"\[([0-9,\s]+)\]", tok)
        if cm and enable_links:
            # render each number as its own hyperlink, keep brackets/commas plain
            inner = cm.group(1)
            # iterate inner numbers
            pieces = re.split(r"([0-9]+)", inner)
            for pc in pieces:
                if pc.strip().isdigit():
                    add_internal_hyperlink(par, pc, f"Ref{pc}")
                else:
                    run = par.add_run(pc)
            continue
        # handle **bold**
        parts = re.split(r"(\*\*.*?\*\*)", tok)
        for p in parts:
            if not p:
                continue
            if p.startswith("**") and p.endswith("**"):
                run = par.add_run(p[2:-2])
                run.bold = True
            else:
                par.add_run(p)


def parse_table_line(line):
    line = line.strip()
    line = line.removeprefix("|")
    line = line.removesuffix("|")
    return [c.strip() for c in line.split("|")]


def is_separator(line):
    return bool(re.match(r"^\s*\|?[\s:|-]+\|?\s*$", line)) and "-" in line


def set_cell_text(cell, text, bold=False):
    """Set cell text, clearing default paragraph first.

    Handles **bold** and ^x^ footnote markers (rendered as superscript).
    """
    cell.text = ""
    p = cell.paragraphs[0]
    # split on bold and footnote markers
    tokens = re.split(r"(\*\*.*?\*\*|\^[a-z]\^)", text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            run = p.add_run(tok[2:-2])
            run.bold = True
        elif re.match(r"^\^[a-z]\^$", tok):
            run = p.add_run(tok[1])  # just the letter, superscripted
            run.font.superscript = True
        else:
            run = p.add_run(tok)
            run.bold = bold
    for r in p.runs:
        r.font.size = Pt(9)
        r.font.name = "Times New Roman"


def split_blocks(lines):
    """Split source lines into markdown blocks separated by blank lines.
    Returns list of (kind, payload):
      kind in {'heading','para','table','figure','ref','abbrev'}
    """
    blocks = []
    cur = []
    for ln in lines:
        if ln.strip():
            cur.append(ln)
        else:
            if cur:
                blocks.append("\n".join(cur))
                cur = []
    if cur:
        blocks.append("\n".join(cur))

    out = []
    for blk in blocks:
        b = blk.strip()
        if b.startswith("## "):
            out.append(("h1", b[3:].strip()))
        elif b.startswith("### "):
            out.append(("h2", b[4:].strip()))
        elif b.startswith("# "):
            out.append(("title", b[2:].strip()))
        elif "|" in b and len(b.splitlines()) > 1 and is_separator(b.splitlines()[1]):
            out.append(("table", b))
        else:
            out.append(("para", b))
    return out


def main():
    md_path = MD
    out_path = PROJ / "paper" / "output" / "doc" / "manuscript.docx"
    if len(sys.argv) > 2:
        md_path = Path(sys.argv[1])
        out_path = Path(sys.argv[2])
    elif len(sys.argv) > 1:
        out_path = Path(sys.argv[1])

    doc = Document()
    # Chinese manuscript -> use SimSun for CJK + Times New Roman for Latin
    is_zh = "zh" in out_path.stem.lower() or "中文" in out_path.name
    cn_font = "SimSun" if is_zh else "Times New Roman"

    # Normal style
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    # set East Asian font (SimSun) so Chinese renders properly
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", cn_font)
    # Heading styles
    for lvl, sz in [(1, 14), (2, 12)]:
        st = doc.styles[f"Heading {lvl}"]
        st.font.name = "Times New Roman"
        st.font.size = Pt(sz)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
        rpr2 = st.element.get_or_add_rPr()
        rpr2.get_or_add_rFonts().set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", cn_font)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    blocks = split_blocks(lines)

    # Special handling: figure legends embed images. Since a figure legend is a
    # markdown block starting with '**Figure N.**', we detect it in the para pass.
    for kind, payload in blocks:
        if kind == "title":
            doc.add_heading(payload, level=0)
        elif kind == "h1":
            doc.add_heading(payload, level=1)
        elif kind == "h2":
            doc.add_heading(payload, level=2)
        elif kind == "table":
            blines = payload.splitlines()
            # find separator row
            sep_idx = None
            for j, bl in enumerate(blines):
                if is_separator(bl):
                    sep_idx = j
                    break
            if sep_idx is None:
                p = doc.add_paragraph()
                add_rich_text(p, payload)
                continue
            header = parse_table_line(blines[0])
            body = [parse_table_line(bl) for bl in blines[sep_idx + 1:] if bl.strip().startswith("|")]
            ncols = max(len(header), *(len(r) for r in body)) if body else len(header)
            table = doc.add_table(rows=1 + len(body), cols=ncols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            # header
            for j in range(ncols):
                set_cell_text(table.cell(0, j), header[j] if j < len(header) else "", bold=True)
            for ri, row in enumerate(body):
                for j in range(ncols):
                    set_cell_text(table.cell(ri + 1, j), row[j] if j < len(row) else "")
            # footnote lines that follow the table (start with ^)
            # (handled as separate para blocks since they start with ^)
            doc.add_paragraph()
        else:  # para (or figure legend)
            # figure marker <!-- FIGURE:N --> -> embed the image here (near text)
            mf = re.match(r"<!-- FIGURE:(\d+) -->", payload)
            if mf:
                fig_no = int(mf.group(1))
                fname = FIG_FILES.get(fig_no)
                fpath = FIG_DIR / fname
                if fpath.exists():
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    try:
                        run.add_picture(str(fpath), width=Inches(5.0))
                    except Exception as e:  # noqa: BLE001 - 单图失败仅警告
                        print(f"  [warn] fig{fig_no}: {e}")
                    # caption under image: pull from the figure legend text
                else:
                    print(f"  [warn] missing {fname}")
                continue
            # figure legend text only (no image; images are placed at markers)
            m = re.match(r"\*\*(Figure|图) (\d+)\.\*\*(.*)", payload)
            if m:
                label = m.group(1)  # 'Figure' or '图'
                fig_no = int(m.group(2))
                p = doc.add_paragraph()
                run = p.add_run(f"{label} {fig_no}.")
                run.bold = True
                add_rich_text(p, m.group(3))
                # 多行图注：续行追加到同一段落（`.*` 不跨行，此前续行被丢弃）
                _, rest = payload.split("\n", 1) if "\n" in payload else (payload, "")
                if rest:
                    add_rich_text(p, " " + " ".join(rest.split()))
            elif re.match(r"^\d{1,2}\.\s", payload):  # reference
                first_line, rest = payload.split("\n", 1) if "\n" in payload else (payload, "")
                mm = re.match(r"^(\d{1,2})\.\s+(.+)$", first_line)
                if mm:
                    p = doc.add_paragraph()
                    run = p.add_run(mm.group(1) + ". ")
                    add_rich_text(p, mm.group(2))
                    if rest:
                        add_rich_text(p, " " + " ".join(rest.split()))
                    # bookmark after all runs so it WRAPS the text (Word-recognized)
                    add_bookmark(p, f"Ref{mm.group(1)}")
                    p.paragraph_format.left_indent = Inches(0.35)
                    p.paragraph_format.first_line_indent = Inches(-0.35)
                else:
                    p = doc.add_paragraph()
                    add_rich_text(p, payload)
            elif re.match(r"^\^[a-z]\^\s*", payload.lstrip()):
                # table footnote marker (^a^ text) -> superscript letter (a) + normal text
                text = " ".join(payload.split())
                m_foot = re.match(r"^(\^)([a-z])(\^)(.*)$", text)
                p = doc.add_paragraph()
                if m_foot:
                    run = p.add_run(m_foot.group(2))  # the letter, superscripted
                    run.font.superscript = True
                    run.font.size = Pt(9)
                    if m_foot.group(4).strip():
                        add_rich_text(p, m_foot.group(4))
                else:
                    add_rich_text(p, text)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.left_indent = Inches(0.2)
            elif payload.lstrip().startswith(">"):
                # blockquote (e.g. Chinese version note) -> italic, indented
                text = " ".join(payload.split())
                text = re.sub(r"^>\s*", "", text)
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.italic = True
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Inches(0.35)
                p.paragraph_format.space_after = Pt(6)
            else:
                # normal paragraph: join all source lines with space
                text = " ".join(payload.split())
                p = doc.add_paragraph()
                add_rich_text(p, text, enable_links=True)
                p.paragraph_format.space_after = Pt(6)

    doc.save(out_path)
    print(f"saved {out_path}")


if __name__ == "__main__":
    main()
